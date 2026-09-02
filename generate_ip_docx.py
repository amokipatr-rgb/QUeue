from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import re

doc = Document()

# Configure default font
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# Title page
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('SMART QUEUE MANAGEMENT SYSTEM')
run.bold = True
run.font.size = Pt(24)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('(SMQSS)')
run.bold = True
run.font.size = Pt(18)

doc.add_paragraph()

inventor = doc.add_paragraph()
inventor.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = inventor.add_run('INVENTION DISCLOSURE DOCUMENT')
run.font.size = Pt(14)

doc.add_paragraph()

author = doc.add_paragraph()
author.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = author.add_run('Inventor: Ogwal Richard')
run.font.size = Pt(12)

author2 = doc.add_paragraph()
author2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = author2.add_run('Student Number: 2300716574')
run.font.size = Pt(12)

advisor = doc.add_paragraph()
advisor.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = advisor.add_run('Advisor: Odongo Steven Eyobu (PhD)')
run.font.size = Pt(12)

university = doc.add_paragraph()
university.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = university.add_run('Makerere University')
run.font.size = Pt(12)

date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = date_p.add_run('August 2026')
run.font.size = Pt(12)

doc.add_page_break()

# Table of Contents
doc.add_heading('Table of Contents', level=1)
toc_items = [
    '1. Title of the Invention',
    '2. Field of the Invention',
    '3. Background of the Invention',
    '4. Summary of the Invention',
    '5. Brief Description of Drawings',
    '6. Detailed Description of Preferred Embodiments',
    '7. Claims (10 Claims)',
    '8. Abstract of the Disclosure',
    '9. Inventor Declaration',
]
for item in toc_items:
    doc.add_paragraph(item, style='List Number')

doc.add_page_break()

# Section 1: Title
doc.add_heading('1. Title of the Invention', level=1)
doc.add_paragraph('SMART QUEUE MANAGEMENT SYSTEM WITH AI-POWERED ANALYTICS, MULTI-PLATFORM KIOSK ARCHITECTURE, AND REAL-TIME VOICE-ENABLED PUBLIC DISPLAYS')

# Section 2: Field
doc.add_heading('2. Field of the Invention', level=1)
doc.add_paragraph('The present invention relates to queue management systems, and more particularly to a smart queue management system integrating artificial intelligence, multi-platform kiosk terminals, real-time public displays with voice announcements, and an administrative dashboard for educational institutions.')

# Section 3: Background
doc.add_heading('3. Background of the Invention', level=1)
doc.add_paragraph('Traditional queue management in educational institutions relies on manual token distribution, paper-based tracking, and verbal announcements. These approaches suffer from:')
problems = [
    'No real-time visibility into queue status for students or administrators',
    'Manual attendance tracking prone to errors and time manipulation',
    'No data-driven insights for service improvement',
    'Fragmented feedback collection with no structured follow-up',
    'Limited accessibility for students with disabilities',
]
for p in problems:
    doc.add_paragraph(p, style='List Bullet')

doc.add_paragraph('Existing commercial queue systems are expensive, proprietary, not designed for educational contexts, and lack AI-powered analytics, voice announcements, and multi-platform kiosk support.')

# Section 4: Summary
doc.add_heading('4. Summary of the Invention', level=1)
doc.add_paragraph('The SMQSS addresses these limitations through:')
features = [
    'Per-day first-free token numbering with gap-aware restart',
    'AI-powered attendance analysis, feedback correlation, and complaint response using LLMs',
    'Real-time public displays with adaptive voice announcements and character spelling',
    'Multi-platform architecture (Electron desktop, web, React Native mobile)',
    'Rate-before-next-token enforcement with QR-coded feedback receipts',
    'Office-hours clamping (8AM-5PM) with dynamic monthly target computation',
    'Three-metric heatmap analytics and tokens-per-hour efficiency metrics',
    'Crash recovery, silent receipt printing, and self-healing database migration',
]
for f in features:
    doc.add_paragraph(f, style='List Bullet')

# Section 5: Drawings
doc.add_heading('5. Brief Description of Drawings', level=1)
drawings = [
    'Figure 1: System Architecture Overview',
    'Figure 2: Token Lifecycle State Diagram',
    'Figure 3: Public Display with Voice Announcement Flow',
    'Figure 4: Officer Dashboard and Attendance Tracking',
    'Figure 5: AI Integration and Analytics Pipeline',
    'Figure 6: Multi-Platform Kiosk Deployment',
    'Figure 7: Database Schema and Migration Flow',
    'Figure 8: Receipt Generation with QR Code',
]
for d in drawings:
    doc.add_paragraph(d, style='List Bullet')

# Section 6: Detailed Description
doc.add_heading('6. Detailed Description of Preferred Embodiments', level=1)

doc.add_heading('6.1 System Architecture', level=2)
doc.add_paragraph('The SMQSS comprises a centralized Flask backend serving REST APIs to multiple frontend platforms. The backend handles token generation, queue management, attendance tracking, feedback collection, AI analytics, voice announcement generation, and administrative functions.')

doc.add_heading('6.2 Token Lifecycle Management', level=2)
doc.add_paragraph('Tokens are generated with per-day first-free numbering. The system queries existing tokens for the current date, extracts numeric suffixes, identifies the first unused number in the sequence, and assigns the token with an office code prefix and zero-padded numeric suffix. A unique constraint on (token_number, token_date) prevents duplicates. Queue reset operations atomically expire waiting tokens, delete expired/skipped tokens, and return the predicted next token number by identifying the first gap in the used number sequence.')

doc.add_heading('6.3 AI Integration', level=2)
doc.add_paragraph('The system integrates a large language model (LLM) for three primary functions: (1) attendance analysis generating natural-language reports with per-officer observations; (2) feedback-officer correlation analysis identifying patterns and improvement recommendations; and (3) complaint response polishing with selectable tone parameters. Monthly grades are computed based on actual working days, with the target dynamically calculated as 540 minutes multiplied by weekdays in the month.')

doc.add_heading('6.4 Public Display and Voice Announcements', level=2)
doc.add_paragraph('The public display receives queue data at regular polling intervals. An adaptive token preview algorithm adjusts the number of preview tokens based on active office count: 3+ offices show 1 next per office; 2 offices show 2 next per office; 1 office shows 3 next tokens. Voice announcements spell token characters individually, concatenate batches into grammatically correct sentences, and deduplicate using composite keys with configurable cooldown.')

doc.add_heading('6.5 Multi-Platform Kiosk Architecture', level=2)
doc.add_paragraph('The system supports three kiosk platforms: Electron desktop (fullscreen, frameless, crash recovery, silent printing), web-based (virtual keyboard adaptation), and React Native/Expo mobile. All platforms communicate with the centralized server through REST API endpoints. The Electron implementation includes a 30-attempt load retry, 30-second watchdog timer, and automatic window recreation on close events.')

doc.add_heading('6.6 Attendance and Analytics', level=2)
doc.add_paragraph('Attendance is calculated using first-login/last-logout with office-hours clamping (8AM-5PM). Multiple daily sessions are merged by taking min(login) and max(logout). Four distinct time metrics are computed: turnaround, service, queue wait, and call response. Three-metric heatmap analytics provide hourly breakdowns of token creation, average wait, and officer presence. Tokens-per-hour efficiency equals daily served tokens multiplied by 60 divided by daily logged-in minutes.')

doc.add_heading('6.7 Security and Infrastructure', level=2)
doc.add_paragraph('Route protection uses token-based URL paths with decoy redirects for unauthorized access attempts. Officer tokens persist to disk; admin and feedback tokens are ephemeral per server start. The database auto-migration system creates tables, adds columns, creates indexes, drops/recreates constraints, backfills data, and seeds defaults on every startup. Geographic IP resolution distinguishes private addresses (Local Network) from public addresses resolved to city, region, country, and GPS coordinates.')

# Section 7: Claims
doc.add_heading('7. Claims', level=1)

claims = [
    {
        'title': 'Claim 1',
        'text': 'A computer-implemented Smart Queue Management System (SMQSS) for managing service queues and service delivery, comprising a centralized server, one or more user terminals, one or more service-officer interfaces, one or more public queue-display terminals, and a feedback interface, wherein the centralized server is configured to receive a service request from a person seeking service; capture identifying information of the person; generate and associate a queue token with the service request; manage the queue token from issuance through service completion; determine and control service order according to queue and service-priority rules; communicate queue status and service-call information in real time; and maintain a service transaction linking the person, queue token, service office, service officer, and service outcome.'
    },
    {
        'title': 'Claim 2',
        'text': 'The system of claim 1, wherein the queue management mechanism is configured to determine availability of a service office and/or service officer before allocating a queue token, capture the name of the person requesting service, associate the person\'s name with the allocated queue token and requested service, and, when the token is called, generate and communicate a service notification identifying the person by name together with the corresponding queue token and service office, thereby linking the person\'s identity, service request, token, service capacity, and service call within the same queue transaction.'
    },
    {
        'title': 'Claim 3',
        'text': 'The system of claim 1, wherein the token management mechanism is configured to generate and assign queue tokens to persons requesting services, associate each token with the corresponding person and service transaction, maintain the token through defined service states comprising waiting, called, serving, completed, cancelled, expired, or skipped, apply configurable service-priority rules to determine service order, and perform controlled queue-reset operations for managing active queue transactions.'
    },
    {
        'title': 'Claim 4',
        'text': 'The system of claim 1, wherein the feedback mechanism is configured to associate the person\'s identity and queue token with a corresponding completed service transaction and enable the person to rate service delivery and submit suggestions, comments, complaints, or recommendations, thereby providing a transaction-linked service evaluation and digital suggestion-and-complaint mechanism, and wherein the system identifies an outstanding unrated completed transaction and controls issuance of a subsequent queue token until the required feedback has been submitted.'
    },
    {
        'title': 'Claim 5',
        'text': 'The system of claim 1, further comprising a receipt and feedback-routing mechanism configured to generate a printed or electronic receipt containing the queue token and a machine-readable code providing access to a corresponding feedback interface, wherein the receipt or code identifies an originating kiosk or service transaction and the originating identifier is preserved throughout the feedback process to return the person to the corresponding kiosk interface following feedback submission.'
    },
    {
        'title': 'Claim 6',
        'text': 'The system of claim 1, further comprising a real-time public display and voice-announcement mechanism configured to dynamically determine and adjust a number of forthcoming queue tokens displayed according to a number of active service offices, generate contextual voice announcements corresponding to queue events, identify a called person by name together with the person\'s queue token and service office, convert token identifiers into individually spoken characters, combine multiple queue announcements into a structured audio message, suppress repeated announcements of the same queue event, and sequentially process announcements to prevent overlapping audio.'
    },
    {
        'title': 'Claim 7',
        'text': 'The system of claim 1, further comprising an artificial-intelligence analytics mechanism configured to receive and correlate attendance, queue, service, feedback, and complaint data to generate service-provider-specific performance observations, attendance analysis, service-efficiency analysis, relationships between service delivery and user feedback, service-improvement recommendations, or assisted responses to complaints; and a multilingual localization mechanism configured to permit persons to interact with the system in a selected language from a plurality of national, official, local, institutional, and community languages, including Ugandan languages, wherein queue instructions, token information, service-status information, feedback functions, notifications, public-display information, voice announcements, and receipts are provided in the selected language while maintaining the underlying queue and service transaction.'
    },
    {
        'title': 'Claim 8',
        'text': 'The system of claim 1, further comprising a multi-platform architecture supporting desktop kiosk, web, and mobile interfaces communicating with the centralized server through application programming interfaces, wherein the architecture further provides automatic recovery of failed or unresponsive kiosk and public-display interfaces and automated database recovery and migration for maintaining operational data and system functions required for continued queue-management and service-delivery operation.'
    },
]

for claim in claims:
    doc.add_heading(claim['title'], level=2)
    doc.add_paragraph(claim['text'])

# Section 8: Abstract
doc.add_heading('8. Abstract of the Disclosure', level=1)
doc.add_paragraph('A smart queue management system comprising an AI-powered Flask backend, multi-platform kiosk terminals (Electron desktop, web, Android mobile), real-time public displays with voice announcements, and an administrative dashboard. The system implements per-day first-free token numbering with gap-aware restart, adaptive display with dynamic preview counts, voice announcements with token character spelling and batch concatenation, AI-powered attendance analysis and feedback correlation, rate-before-next-token enforcement, and kiosk-type-aware feedback redirect tracking. The system supports crash recovery, silent receipt printing, geographic IP tracking, and self-healing database migration.')

# Section 9: Declaration
doc.add_heading('9. Inventor Declaration', level=1)
doc.add_paragraph('I, Ogwal Richard, hereby declare that I am the original inventor of the Smart Queue Management System (SMQSS) described in this document, developed under the supervision of Odongo Steven Eyobu (PhD) at Makerere University. All claims herein are based on original research and development conducted between January 2025 and August 2026.')

doc.add_paragraph()
doc.add_paragraph('Inventor: Ogwal Richard')
doc.add_paragraph('Student Number: 2300716574')
doc.add_paragraph('Signature: ________________________')
doc.add_paragraph('Date: ________________________')
doc.add_paragraph()
doc.add_paragraph('Advisor: Odongo Steven Eyobu (PhD)')
doc.add_paragraph('Signature: ________________________')
doc.add_paragraph('Date: ________________________')
doc.add_paragraph()
doc.add_paragraph('Document prepared for intellectual property protection under the Uganda Industrial Property Act, 2003 and the ARIPO Harare Protocol on Patents.')

# Save
doc.save('SMQSS_IP_PATENT.docx')
print('DOCX generated successfully: SMQSS_IP_PATENT.docx')
